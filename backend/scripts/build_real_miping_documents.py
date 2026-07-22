"""Build realistic initial and remediated document sets for Miping acceptance."""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "artifacts" / "real-miping-fixtures"
INITIAL_GAP = OUTPUT / "initial" / "01-密评准备与差距分析"
INITIAL_FIELD = OUTPUT / "initial" / "02-密码应用现场评估"
REMEDIATED = OUTPUT / "remediated"
LIBRARY_PATH = ROOT / "reference" / "compliance" / "miping_document_controls.yaml"
DOCX_FONT = "Arial Unicode MS"

GAP_KEYS = {
    "system_crypto_scope",
    "crypto_application_plan",
    "crypto_algorithm_protocol_inventory",
    "crypto_product_inventory",
    "key_management_policy",
    "crypto_management_policy",
    "crypto_personnel_records",
    "crypto_operation_records",
    "crypto_incident_plan",
}
FIELD_KEYS = {
    "physical_environment_evidence",
    "network_communication_evidence",
    "device_computing_evidence",
    "application_data_evidence",
}
INITIAL_OMISSIONS = {
    "crypto_application_plan": {"MIP-PLAN-005"},
    "crypto_algorithm_protocol_inventory": {"MIP-ALG-003"},
    "key_management_policy": {"MIP-KEY-005"},
    "crypto_product_inventory": {"MIP-PROD-004"},
    "network_communication_evidence": {"MIP-NET-004"},
    "application_data_evidence": {"MIP-APP-006"},
}


def file_name(key: str, name: str, remediated: bool) -> str:
    suffix = "-整改复测版" if remediated else "-初检版"
    return f"{name}{suffix}.docx"


def evidence_text(document_name: str, control: dict, point: dict, index: int) -> str:
    keywords = "、".join(point.get("evidence_keywords") or [])
    return (
        f"{index}. {point['text']}。示例企业已将本要求纳入《{document_name}》执行范围，"
        f"由密码应用负责人统筹，系统管理员实施，密码审计员独立复核。具体实施覆盖{keywords}。"
        "执行前应确认保护对象、业务场景、算法协议、密码产品服务及密钥用途；执行后保存审批单、"
        "配置截图、设备或服务日志、验证结果和异常处置记录。每季度抽样核验一次，系统、网络、"
        "重要数据或密码配置发生变化时立即重新评估；发现偏差后十个工作日内完成整改和复测。"
    )


def active_controls(spec: dict, key: str, remediated: bool) -> list[dict]:
    omitted = set() if remediated else INITIAL_OMISSIONS.get(key, set())
    return [control for control in spec["controls"] if control["id"] not in omitted]


def configure_docx(document: Document, title: str, remediated: bool) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    normal = document.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in (("Title", 22, "0B2545"), ("Heading 1", 15, "2E74B5"), ("Heading 2", 12, "1F4D78")):
        style = document.styles[style_name]
        style.font.name = DOCX_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header.paragraphs[0]
    header.text = f"CERTIPROOF 密评真实材料验收 | {title}"
    footer = section.footer.paragraphs[0]
    footer.text = "仅用于企业内部密评自查自动化验收，不代替具备资质机构出具的正式密评结论。"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading(title, 0)
    document.add_paragraph("整改复测材料" if remediated else "初次检查材料（故意保留部分缺口）")
    metadata = document.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata.autofit = False
    values = [
        ("文件状态", "整改后正式版" if remediated else "初检待改进版"),
        ("适用系统", "示例企业核心业务系统"),
        ("责任部门", "密码应用管理办公室"),
        ("版本日期", "V2.0 / 2026-07-22" if remediated else "V1.0 / 2026-07-01"),
    ]
    for row, values_row in zip(metadata.rows, values):
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(11)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[0].text, row.cells[1].text = values_row
        row.cells[0].paragraphs[0].runs[0].bold = True


def apply_font(document: Document) -> None:
    containers = list(document.paragraphs)
    for table in document.tables:
        containers.extend(paragraph for row in table.rows for cell in row.cells for paragraph in cell.paragraphs)
    for section in document.sections:
        containers.extend(section.header.paragraphs)
        containers.extend(section.footer.paragraphs)
    for paragraph in containers:
        paragraph.paragraph_format.keep_together = True
        for run in paragraph.runs:
            run.font.name = DOCX_FONT
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT)


def build_docx(path: Path, key: str, spec: dict, remediated: bool) -> None:
    document = Document()
    configure_docx(document, spec["name"], remediated)
    document.add_heading("1. 目的、范围与责任", level=1)
    document.add_paragraph(
        "本文件适用于示例企业核心业务系统及其物理环境、网络通信、设备计算、应用和数据。"
        "密码应用负责人批准要求，密码管理员负责实施，业务和系统责任人提供对象信息，审计员复核证据。"
    )
    document.add_heading("2. 控制要求与执行证据", level=1)
    index = 1
    for control in active_controls(spec, key, remediated):
        document.add_heading(f"2.{index} {control['title']}（{control['id']}）", level=2)
        for point in control["required_points"]:
            document.add_paragraph(evidence_text(spec["name"], control, point, index))
            index += 1
    if not remediated and INITIAL_OMISSIONS.get(key):
        document.add_heading("3. 当前待完善事项", level=1)
        document.add_paragraph("本版本仍有专项控制尚未形成完整设计、配置或运行证据，应完成整改后重新提交全量材料复测。")
    document.add_heading("4. 检查、变更与留存", level=1)
    document.add_paragraph(
        "责任部门每季度检查制度、配置和运行记录的一致性。涉及密码算法、产品、服务、证书、密钥、"
        "重要数据或系统边界变化时发起重新评估。原始检查记录、自查报告和问题处置记录至少保存六年。"
    )
    document.core_properties.title = spec["name"]
    document.core_properties.subject = "CertiProof 密评真实材料闭环验收"
    document.core_properties.author = "示例企业密码应用管理办公室"
    apply_font(document)
    document.save(path)


def make_zip(path: Path, files: list[Path]) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for file_path in files:
            archive.write(file_path, file_path.name)


def main() -> None:
    library = yaml.safe_load(LIBRARY_PATH.read_text(encoding="utf-8"))
    for directory in (INITIAL_GAP, INITIAL_FIELD, REMEDIATED):
        directory.mkdir(parents=True, exist_ok=True)
        for old in directory.iterdir():
            if old.is_file():
                old.unlink()

    initial_gap_files: list[Path] = []
    initial_field_files: list[Path] = []
    remediated_files: list[Path] = []
    manifest = {"standard_version": library["version"], "initial": [], "remediated": []}
    for key, spec in library["documents"].items():
        initial_dir = INITIAL_GAP if key in GAP_KEYS else INITIAL_FIELD
        initial_path = initial_dir / file_name(key, spec["name"], False)
        remediated_path = REMEDIATED / file_name(key, spec["name"], True)
        build_docx(initial_path, key, spec, False)
        build_docx(remediated_path, key, spec, True)
        (initial_gap_files if key in GAP_KEYS else initial_field_files).append(initial_path)
        remediated_files.append(remediated_path)
        manifest["initial"].append({"key": key, "file": str(initial_path), "omitted_controls": sorted(INITIAL_OMISSIONS.get(key, set()))})
        manifest["remediated"].append({"key": key, "file": str(remediated_path)})

    make_zip(OUTPUT / "密评初检-准备与差距分析材料.zip", initial_gap_files)
    make_zip(OUTPUT / "密评初检-现场证据材料.zip", initial_field_files)
    make_zip(OUTPUT / "密评整改复测-全量材料.zip", remediated_files)
    (OUTPUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "output": str(OUTPUT),
        "gap_files": len(initial_gap_files),
        "field_files": len(initial_field_files),
        "remediated_files": len(remediated_files),
        "deliberate_gaps": sum(len(value) for value in INITIAL_OMISSIONS.values()),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
