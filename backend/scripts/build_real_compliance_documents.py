"""Build realistic mixed-format documents for the full compliance-flow acceptance test."""

from __future__ import annotations

import json
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

try:
    import yaml
except ModuleNotFoundError:
    sys.path.append(str(Path.home() / "Library/Python/3.9/lib/python/site-packages"))
    import yaml


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "artifacts" / "real-compliance-fixtures"
INITIAL = OUTPUT / "initial"
REMEDIATED = OUTPUT / "remediated"
SUPPORT = OUTPUT / "support"
COMPARISON = OUTPUT / "整改复测对照材料"
INITIAL_ISSUES = COMPARISON / "01-初次检测（完整材料，含问题）"
RETEST_DOCUMENTS = COMPARISON / "02-复测（已修改）"
CONTROL_LIBRARY = ROOT / "reference" / "compliance" / "document_controls.yaml"
CHINESE_FONT = "/System/Library/Fonts/STHeiti Light.ttc"
DOCX_FONT = "Noto Sans CJK SC"

FORMATS = {
    "security_management_policy": "docx",
    "security_org_setup": "docx",
    "personnel_security_policy": "docx",
    "secure_construction_policy": "pdf",
    "security_operations_policy": "docx",
    "incident_response_plan": "scan_pdf",
    "incident_management_policy": "docx",
    "security_audit_policy": "docx",
    "system_security_plan": "pdf",
    "security_strategy": "docx",
}

FILE_NAMES = {
    "security_management_policy": "信息安全管理制度V2.1.docx",
    "security_org_setup": "关于成立网络安全领导小组的决定V2.docx",
    "personnel_security_policy": "人员安全管理制度V1.0.docx",
    "secure_construction_policy": "安全建设管理制度（正式发布版）.pdf",
    "security_operations_policy": "安全运维管理制度V3.2.docx",
    "incident_response_plan": "信息安全事件应急预案（盖章扫描件）.pdf",
    "incident_management_policy": "安全事件管理制度V1.0.docx",
    "security_audit_policy": "安全审计管理制度2026版.docx",
    "system_security_plan": "核心业务系统安全方案.pdf",
    "security_strategy": "网络与信息安全总体策略V2.docx",
}

# Initial versions deliberately omit these controls so remediation has real work to do.
OMISSIONS = {
    "personnel_security_policy": {"DOC-PSP-004"},
    "incident_management_policy": {"DOC-IMP-003"},
}
NATIVE_RECOVERY_DOCUMENTS = {"incident_response_plan"}

OWNERS = {
    "security_management_policy": "信息安全管理部",
    "security_org_setup": "网络安全领导小组办公室",
    "personnel_security_policy": "人力资源部与信息安全管理部",
    "secure_construction_policy": "项目管理办公室与研发中心",
    "security_operations_policy": "基础设施运维部",
    "incident_response_plan": "信息安全事件应急指挥组",
    "incident_management_policy": "安全运营中心",
    "security_audit_policy": "安全审计岗",
    "system_security_plan": "系统建设项目组",
    "security_strategy": "信息安全委员会",
}


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])
    paragraph.add_run(" 页")


def configure_docx(document: Document, title: str, code: str) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = DOCX_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after, color in (
        ("Title", 23, 0, 6, "0B2545"),
        ("Heading 1", 16, 16, 8, "2E74B5"),
        ("Heading 2", 13, 12, 6, "2E74B5"),
        ("Heading 3", 12, 8, 4, "1F4D78"),
    ):
        style = styles[name]
        style.font.name = DOCX_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = f"CERTIPROOF 真实材料验收  |  {title}"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = DOCX_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(103, 119, 132)
    add_page_number(section.footer.paragraphs[0])

    title_paragraph = document.add_paragraph(style="Title")
    title_paragraph.add_run(title).bold = True
    subtitle = document.add_paragraph("企业等保三级自查制度文件", style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(16)

    table = document.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    metadata = [
        ("文件编号", code),
        ("版本/状态", "V2.1 / 正式发布"),
        ("发布/生效", "2026-07-01 / 2026-07-15"),
        ("批准/归口", "总经理 / 信息安全管理部"),
    ]
    for row, (label, value) in zip(table.rows, metadata):
        set_cell_width(row.cells[0], 2160)
        set_cell_width(row.cells[1], 7200)
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "F2F4F7")
        row.cells[0].paragraphs[0].runs[0].bold = True


def point_paragraph(document_name: str, owner: str, point: dict, index: int) -> str:
    keywords = "、".join(point.get("evidence_keywords") or [])
    return (
        f"第{index}条  {owner}负责落实“{point['text']}”。本条适用于公司总部、分支机构、正式员工、"
        f"外包人员以及纳入本次测评的业务系统和信息资产，执行内容包括{keywords}。责任人应在相关事项发生前完成"
        f"审批或登记，并在事项结束后两个工作日内将申请单、审批记录、执行日志和复核结论归档至安全管理平台。"
        f"部门负责人每季度检查一次执行情况，发现偏差应在十个工作日内完成纠正；记录至少保存三年，审计岗可按"
        f"文件、责任人、日期和业务对象追溯。该要求由{document_name}发布之日起执行。"
    )


def document_sections(document_key: str, spec: dict, remediated: bool) -> list[dict]:
    omitted = set() if remediated else OMISSIONS.get(document_key, set())
    sections = []
    for control in spec["controls"]:
        if control["id"] in omitted:
            continue
        sections.append({
            "id": control["id"],
            "title": control["title"],
            "points": control["required_points"],
        })
    return sections


def add_common_content_docx(document: Document, key: str, spec: dict, remediated: bool) -> None:
    owner = OWNERS[key]
    document.add_heading("1. 目的与适用范围", level=1)
    document.add_paragraph(
        f"本文件用于建立可执行、可检查、可追溯的信息安全管理要求。适用于公司各部门、全体员工、外包人员、"
        f"供应商以及生产、测试和办公环境中的信息系统、网络、数据和设备。归口责任部门为{owner}。"
    )
    document.add_heading("2. 角色与职责", level=1)
    roles = document.add_table(rows=1, cols=3)
    roles.alignment = WD_TABLE_ALIGNMENT.CENTER
    roles.autofit = False
    for cell, width in zip(roles.rows[0].cells, (2160, 2880, 4320)):
        set_cell_width(cell, width)
        set_cell_shading(cell, "E8EEF5")
    roles.rows[0].cells[0].text = "角色"
    roles.rows[0].cells[1].text = "责任主体"
    roles.rows[0].cells[2].text = "主要职责"
    for role, subject, duty in (
        ("批准人", "总经理", "批准制度、重大例外与资源投入"),
        ("归口部门", owner, "维护制度、组织执行、检查整改与保留记录"),
        ("执行部门", "各业务与技术部门", "落实条款、提交证据并对偏差负责"),
        ("独立监督", "安全审计岗", "按季度抽查并向信息安全委员会报告"),
    ):
        cells = roles.add_row().cells
        for cell, width in zip(cells, (2160, 2880, 4320)):
            set_cell_width(cell, width)
        for cell, value in zip(cells, (role, subject, duty)):
            cell.text = value

    document.add_heading("3. 管理要求", level=1)
    point_index = 1
    for section in document_sections(key, spec, remediated):
        document.add_heading(f"3.{point_index} {section['title']}", level=2)
        for point in section["points"]:
            document.add_paragraph(point_paragraph(spec["name"], owner, point, point_index))
            point_index += 1

    if not remediated and key in OMISSIONS:
        document.add_paragraph(
            "说明：本版本尚未完成全部专项管理条款的发布，缺失内容将在下一版本补充。",
            style="Intense Quote",
        )

    document.add_heading("4. 执行、检查与记录", level=1)
    document.add_paragraph(
        f"{owner}每季度首月组织制度执行检查，检查范围覆盖责任落实、审批流程、记录完整性和问题关闭情况。"
        "检查结果形成《信息安全制度执行检查表》，严重偏差二十四小时内上报，普通偏差十个工作日内整改。"
    )
    records = document.add_table(rows=1, cols=4)
    records.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("记录名称", "责任部门", "保存期限", "复核频率")
    for cell, value in zip(records.rows[0].cells, headers):
        cell.text = value
        set_cell_shading(cell, "E8EEF5")
    for values in (
        ("审批与授权记录", owner, "不少于3年", "每季度"),
        ("执行日志与检查表", "各执行部门", "不少于3年", "每季度"),
        ("偏差与整改记录", owner, "关闭后3年", "每月跟踪"),
    ):
        cells = records.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value

    document.add_heading("5. 审批、例外与修订", level=1)
    document.add_paragraph(
        "本文件经归口部门会签、信息安全委员会审议并由总经理批准后发布生效。任何例外或豁免必须提交书面申请，"
        "说明业务理由、影响范围、补偿措施、责任人和失效日期，经风险评估及批准后方可执行。归口部门每年至少评审"
        "一次；法律法规、组织架构、业务系统或风险发生重大变化时应即时修订，版本变更和作废记录永久保留。"
    )


def add_org_chart(document: Document) -> None:
    image_path = SUPPORT / "security-organization-chart.png"
    image = Image.new("RGB", (1200, 560), "white")
    draw = ImageDraw.Draw(image)
    font = ImageFont.truetype(CHINESE_FONT, 34)
    small = ImageFont.truetype(CHINESE_FONT, 26)
    boxes = [
        (390, 35, 810, 120, "信息安全委员会"),
        (390, 225, 810, 310, "信息安全管理部"),
        (50, 420, 340, 510, "安全运营中心"),
        (455, 420, 745, 510, "系统与运维组"),
        (860, 420, 1150, 510, "安全审计岗"),
    ]
    for x1, y1, x2, y2, label in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=12, outline="#1f6d8a", width=4, fill="#edf7fa")
        draw.text(((x1 + x2) / 2, (y1 + y2) / 2), label, font=font if y1 < 300 else small, fill="#123647", anchor="mm")
    for start, end in (((600, 120), (600, 225)), ((600, 310), (195, 420)), ((600, 310), (600, 420)), ((600, 310), (1005, 420))):
        draw.line((*start, *end), fill="#34a4bd", width=4)
    image.save(image_path)
    document.add_heading("附图：安全管理组织关系", level=1)
    document.add_picture(str(image_path), width=Inches(6.4))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_docx_fonts(document: Document) -> None:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(paragraph for row in table.rows for cell in row.cells for paragraph in cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = DOCX_FONT
            run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_FONT)


def build_docx(path: Path, key: str, spec: dict, remediated: bool) -> None:
    document = Document()
    configure_docx(document, spec["name"], f"CP-{key.upper()[:8]}-2026")
    add_common_content_docx(document, key, spec, remediated)
    if key == "security_org_setup":
        add_org_chart(document)
    document.core_properties.title = spec["name"]
    document.core_properties.subject = "CertiProof 等保文档合规真实材料验收"
    document.core_properties.author = "示例企业信息安全管理部"
    document.core_properties.comments = "仅用于 CertiProof 自动化验收，不代表正式合规结论。"
    apply_docx_fonts(document)
    document.save(path)


def pdf_styles():
    if "STHeiti" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("STHeiti", CHINESE_FONT, subfontIndex=1))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("ChineseTitle", parent=styles["Title"], fontName="STHeiti", fontSize=22, leading=30, textColor=colors.HexColor("#0B2545"), alignment=TA_CENTER, spaceAfter=18),
        "h1": ParagraphStyle("ChineseH1", parent=styles["Heading1"], fontName="STHeiti", fontSize=15, leading=21, textColor=colors.HexColor("#2E74B5"), spaceBefore=12, spaceAfter=7),
        "h2": ParagraphStyle("ChineseH2", parent=styles["Heading2"], fontName="STHeiti", fontSize=12, leading=18, textColor=colors.HexColor("#1F4D78"), spaceBefore=9, spaceAfter=5),
        "body": ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName="STHeiti", fontSize=10.5, leading=17, textColor=colors.HexColor("#172B3A"), spaceAfter=7),
        "small": ParagraphStyle("ChineseSmall", parent=styles["BodyText"], fontName="STHeiti", fontSize=8.5, leading=12, textColor=colors.HexColor("#607080")),
    }


def build_text_pdf(path: Path, key: str, spec: dict, remediated: bool) -> None:
    styles = pdf_styles()

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("STHeiti", 8)
        canvas.setFillColor(colors.HexColor("#607080"))
        canvas.drawString(0.85 * inch, 0.48 * inch, f"CERTIPROOF 真实材料验收 | {spec['name']}")
        canvas.drawRightString(7.65 * inch, 0.48 * inch, f"第 {doc.page} 页")
        canvas.restoreState()

    story = [
        Paragraph(spec["name"], styles["title"]),
        Paragraph("企业等保三级自查制度文件 | V2.1 正式发布 | 2026-07-15 生效", styles["small"]),
        Spacer(1, 12),
        Table(
            [["文件编号", f"CP-{key.upper()[:8]}-2026"], ["批准人", "总经理"], ["归口部门", OWNERS[key]], ["保存期限", "现行版本长期保存，执行记录不少于三年"]],
            colWidths=[1.35 * inch, 5.1 * inch],
            style=TableStyle([
                ("FONT", (0, 0), (-1, -1), "STHeiti", 9),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8C4CC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]),
        ),
        Paragraph("1. 目的与适用范围", styles["h1"]),
        Paragraph(f"本文件适用于公司总部、分支机构、员工、外包人员、供应商以及生产、测试和办公环境中的信息系统、网络、数据和设备。归口责任部门为{OWNERS[key]}。", styles["body"]),
        Paragraph("2. 管理要求", styles["h1"]),
    ]
    number = 1
    for section in document_sections(key, spec, remediated):
        story.append(Paragraph(f"2.{number} {section['title']}", styles["h2"]))
        for point in section["points"]:
            story.append(Paragraph(point_paragraph(spec["name"], OWNERS[key], point, number), styles["body"]))
            number += 1
    story.extend([
        Paragraph("3. 检查、记录与改进", styles["h1"]),
        Paragraph(f"{OWNERS[key]}每季度组织一次检查并形成检查表。审批单、执行日志、配置清单、复核结论和整改记录至少保存三年；严重偏差二十四小时内上报，普通偏差十个工作日内完成整改。", styles["body"]),
        Paragraph("4. 审批、例外与版本维护", styles["h1"]),
        Paragraph("文件经归口部门会签、信息安全委员会审议并由总经理批准后发布。例外必须书面说明影响、补偿措施、责任人与失效日期。每年至少评审一次，重大变化时即时修订并保留完整版本记录。", styles["body"]),
    ])
    SimpleDocTemplate(str(path), pagesize=letter, leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=0.8 * inch, title=spec["name"], author="示例企业信息安全管理部").build(story, onFirstPage=footer, onLaterPages=footer)


def wrap_chinese(draw, text: str, font, max_width: int) -> list[str]:
    lines, current = [], ""
    for char in text:
        candidate = current + char
        if draw.textlength(candidate, font=font) > max_width and current:
            lines.append(current)
            current = char
        else:
            current = candidate
    if current:
        lines.append(current)
    return lines


def build_scanned_pdf(path: Path, key: str, spec: dict, remediated: bool) -> None:
    width, height = 1654, 2339
    title_font = ImageFont.truetype(CHINESE_FONT, 52)
    heading_font = ImageFont.truetype(CHINESE_FONT, 34)
    body_font = ImageFont.truetype(CHINESE_FONT, 25)
    small_font = ImageFont.truetype(CHINESE_FONT, 20)
    paragraphs = []
    for index, section in enumerate(document_sections(key, spec, remediated), 1):
        paragraphs.append(("heading", f"{index}. {section['title']}"))
        for point in section["points"]:
            paragraphs.append(("body", point_paragraph(spec["name"], OWNERS[key], point, index)))
    paragraphs.extend([
        ("heading", "执行检查与记录"),
        ("body", f"{OWNERS[key]}每季度组织检查，审批单、处置记录、演练记录、复盘报告和整改记录至少保存三年。严重事件二十四小时内上报，普通偏差十个工作日内完成整改。"),
        ("heading", "审批与修订"),
        ("body", "本预案经信息安全委员会审议并由总经理批准后发布；每年至少演练和评审一次，重大事件后五个工作日内复盘并更新预案。"),
    ])
    pages = []
    page = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(page)
    y = 150
    draw.text((width / 2, y), spec["name"], font=title_font, fill="#102f43", anchor="ma")
    y += 90
    draw.text((width / 2, y), "V2.1 正式发布 | 2026-07-15 生效 | 批准人：总经理", font=small_font, fill="#536b79", anchor="ma")
    y += 100
    for kind, text in paragraphs:
        font = heading_font if kind == "heading" else body_font
        fill = "#1f6380" if kind == "heading" else "#172b3a"
        spacing = 50 if kind == "heading" else 39
        lines = wrap_chinese(draw, text, font, width - 260)
        required = len(lines) * spacing + (35 if kind == "heading" else 22)
        if y + required > height - 150:
            draw.text((width - 130, height - 90), f"第 {len(pages) + 1} 页", font=small_font, fill="#71818b", anchor="ra")
            pages.append(page)
            page = Image.new("RGB", (width, height), "white")
            draw = ImageDraw.Draw(page)
            y = 125
            draw.text((130, 60), f"CERTIPROOF 真实材料验收 | {spec['name']}", font=small_font, fill="#71818b")
        for line in lines:
            draw.text((130, y), line, font=font, fill=fill)
            y += spacing
        y += 35 if kind == "heading" else 22
    draw.text((width - 130, height - 90), f"第 {len(pages) + 1} 页", font=small_font, fill="#71818b", anchor="ra")
    pages.append(page)
    pages[0].save(path, "PDF", resolution=150, save_all=True, append_images=pages[1:])


def build_unrelated_doc(path: Path) -> None:
    document = Document()
    configure_docx(document, "员工食堂值班安排", "ADM-CANTEEN-2026")
    document.add_heading("一、本周安排", level=1)
    document.add_paragraph("周一至周五由行政部安排午间值班，负责餐具盘点、就餐秩序和卫生巡查。")
    document.add_heading("二、联系方式", level=1)
    document.add_paragraph("临时调整请联系行政前台。本文件不涉及信息系统、网络安全、数据安全或人员权限管理。")
    apply_docx_fonts(document)
    document.save(path)


def main() -> None:
    for directory in (INITIAL, REMEDIATED, SUPPORT):
        directory.mkdir(parents=True, exist_ok=True)
    shutil.rmtree(COMPARISON, ignore_errors=True)
    INITIAL_ISSUES.mkdir(parents=True)
    RETEST_DOCUMENTS.mkdir(parents=True)
    library = yaml.safe_load(CONTROL_LIBRARY.read_text(encoding="utf-8"))
    manifest = {"library_version": library["version"], "initial": [], "remediated": [], "edge_cases": []}
    for key, spec in library["documents"].items():
        fmt = FORMATS[key]
        initial_path = INITIAL / FILE_NAMES[key]
        if fmt == "docx":
            build_docx(initial_path, key, spec, False)
        elif fmt == "pdf":
            build_text_pdf(initial_path, key, spec, False)
        else:
            build_scanned_pdf(initial_path, key, spec, False)
        manifest["initial"].append({
            "document_key": key,
            "document_name": spec["name"],
            "file": initial_path.name,
            "format": fmt,
            "expected": "partial" if key in OMISSIONS else "pass",
            "omitted_controls": sorted(OMISSIONS.get(key, set())),
        })
        if key in OMISSIONS or key in NATIVE_RECOVERY_DOCUMENTS:
            suffix = ".docx" if key in NATIVE_RECOVERY_DOCUMENTS else initial_path.suffix
            label = "可解析版" if key in NATIVE_RECOVERY_DOCUMENTS else "整改版"
            base_name = spec["name"] if key in NATIVE_RECOVERY_DOCUMENTS else initial_path.stem
            remediation_path = REMEDIATED / f"{base_name}-{label}{suffix}"
            if suffix == ".docx":
                build_docx(remediation_path, key, spec, True)
            elif suffix == ".pdf":
                build_text_pdf(remediation_path, key, spec, True)
            else:
                build_scanned_pdf(remediation_path, key, spec, True)
            manifest["remediated"].append({
                "document_key": key,
                "document_name": spec["name"],
                "file": remediation_path.name,
                "expected": "pass",
                "purpose": "visual_model_recovery" if key in NATIVE_RECOVERY_DOCUMENTS else "content_remediation",
            })

    unrelated = INITIAL / "员工食堂值班安排.docx"
    build_unrelated_doc(unrelated)
    manifest["edge_cases"].append({"file": unrelated.name, "expected": "unclassified_warning"})

    for initial_source in sorted(INITIAL.iterdir()):
        if initial_source.is_file():
            shutil.copy2(initial_source, INITIAL_ISSUES / initial_source.name)

    archive_path = OUTPUT / "初次差距分析材料包.zip"
    with zipfile.ZipFile(archive_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(INITIAL.iterdir()):
            archive.write(path, arcname=f"等保自查材料/{path.name}")
    manifest["batch_archive"] = archive_path.name

    initial_by_key = {item["document_key"]: item for item in manifest["initial"]}
    comparison_rows = []
    for item in manifest["remediated"]:
        initial_item = initial_by_key[item["document_key"]]
        initial_source = INITIAL / initial_item["file"]
        retest_source = REMEDIATED / item["file"]
        shutil.copy2(retest_source, RETEST_DOCUMENTS / retest_source.name)
        issue = (
            "扫描版材料用于验证视觉解析失败后的重新提交"
            if item["purpose"] == "visual_model_recovery"
            else f"缺少必检项：{', '.join(initial_item['omitted_controls'])}"
        )
        comparison_rows.append(
            f"| {item['document_name']} | {initial_source.name} | {issue} | {retest_source.name} | 通过 |"
        )

    instructions = "\n".join([
        "# 整改复测对照材料",
        "",
        "1. 先上传 `01-初次检测（完整材料，含问题）`，执行全部文档合规检查并保留问题结果。",
        "2. 进入整改与复测，上传 `02-复测（已修改）` 中对应文件。",
        "3. 系统应重新分析同类文档，关闭已修复问题，并保留初检与复测对比记录。",
        "",
        "| 文档类型 | 初检文件 | 预期问题 | 复测文件 | 复测预期 |",
        "| --- | --- | --- | --- | --- |",
        *comparison_rows,
        "",
    ])
    (COMPARISON / "使用说明.md").write_text(instructions, encoding="utf-8")

    package_paths = []
    for package_name, source_dir in (
        ("01-初次检测完整材料包.zip", INITIAL_ISSUES),
        ("02-整改后复测文档.zip", RETEST_DOCUMENTS),
    ):
        package_path = OUTPUT / package_name
        with zipfile.ZipFile(package_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in sorted(source_dir.iterdir()):
                archive.write(path, arcname=path.name)
        package_paths.append(package_path.name)

    (OUTPUT / "01-初次检测问题文档.zip").unlink(missing_ok=True)

    comparison_archive = OUTPUT / "整改复测对照材料包.zip"
    with zipfile.ZipFile(comparison_archive, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(COMPARISON.rglob("*")):
            if path.is_file():
                archive.write(path, arcname=path.relative_to(COMPARISON.parent))
    package_paths.append(comparison_archive.name)
    manifest["verification_packages"] = package_paths
    (OUTPUT / "manifest.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({
        "output": str(OUTPUT),
        "initial_files": len(manifest["initial"]),
        "remediated_files": len(manifest["remediated"]),
        "edge_cases": len(manifest["edge_cases"]),
        "archive": str(archive_path),
        "verification_packages": package_paths,
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
