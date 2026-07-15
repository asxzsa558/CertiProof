import io
import tempfile
import zipfile
from pathlib import Path
from types import SimpleNamespace

from docx import Document
import yaml

from app.services.document_control_engine import DocumentControlEngine
from app.services.document_pipeline import _deduplicate, _docx_native, expand_document_upload


LIBRARY = yaml.safe_load((Path(__file__).resolve().parents[2] / "reference" / "compliance" / "document_controls.yaml").read_text(encoding="utf-8"))


with tempfile.TemporaryDirectory() as directory:
    path = Path(directory) / "arbitrary-name-v2.docx"
    document = Document()
    document.sections[0].header.paragraphs[0].text = "受控文件"
    document.sections[0].footer.paragraphs[0].text = "第 1 页"
    document.add_heading("安全事件管理", level=1)
    document.add_paragraph("安全事件应当分级、报告并及时处置。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "责任人"
    table.cell(0, 1).text = "安全管理员"
    table.cell(1, 0).text = "处置要求"
    table.cell(1, 1).text = "记录、复盘和改进"
    document.save(path)

    evidence = SimpleNamespace(id=1, original_name=path.name)
    blocks, images, pages = _docx_native(path, evidence)
    types = {block["type"] for block in blocks}
    assert {"heading", "text", "table", "header", "footer"} <= types
    assert not images
    assert pages == 1

    duplicate = dict(blocks[0], block_id="duplicate", source="vision")
    merged = _deduplicate([duplicate, *blocks])
    assert len(merged) == len(blocks)
    assert merged[0]["source"] == "native"

    analysis = DocumentControlEngine(LIBRARY).analyze_blocks(blocks, "安全事件管理制度")
    assert analysis["status"] in {"pass", "partial", "fail"}
    evidence_items = [
        item
        for control in analysis["controls"]
        for point in control["points"]
        for item in point["evidence"]
    ]
    assert evidence_items
    assert all(item["document_file_id"] == 1 for item in evidence_items)
    assert all(item["file_name"] == path.name for item in evidence_items)
    report_point = analysis["controls"][0]["points"][0]
    assert report_point["evidence"], "同一内容块内的“事件…报告”应召回“事件报告”检查点"

    archive_buffer = io.BytesIO()
    with zipfile.ZipFile(archive_buffer, "w") as archive:
        archive.writestr("制度/安全事件管理制度V2.txt", "安全事件报告和处置")
        archive.writestr("__MACOSX/._metadata", "ignored")
        archive.writestr("说明.exe", "ignored")
    documents, skipped = expand_document_upload("制度包.zip", archive_buffer.getvalue())
    assert [name for name, _ in documents] == ["制度/安全事件管理制度V2.txt"]
    assert skipped == ["说明.exe"]

print("document pipeline checks passed")
