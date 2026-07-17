"""Document extraction and asynchronous compliance analysis."""

from __future__ import annotations

import asyncio
import base64
import contextlib
import hashlib
import io
import logging
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import datetime, timedelta
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any

import httpx
from sqlalchemy import delete, or_, select, update
from sqlalchemy.exc import SQLAlchemyError
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import AsyncSessionLocal
from app.models.assessment import Assessment, PhaseInstance, TaskInstance
from app.models.document_knowledge import (
    DocumentAnalysisRun,
    DocumentBlock,
    DocumentControlResult,
    DocumentEvidenceLink,
    DocumentFile,
    DocumentRunFile,
)
from app.services.document_control_engine import DocumentControlEngine
from app.services.file_storage import file_storage
from app.services.knowledge_graph import knowledge_graph

logger = logging.getLogger(__name__)
DOCUMENT_SOURCE = "document_control_analysis"
BATCH_DOCUMENT_SOURCE = "document_batch_classification"
SUPPORTED_SUFFIXES = {".docx", ".pdf", ".txt", ".md", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
ARCHIVE_SUFFIXES = {".zip", ".rar", ".7z"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
ANALYSIS_MODES = {"standard", "deep"}
MAX_BATCH_FILES = 100
MAX_BATCH_UNCOMPRESSED = 300 * 1024 * 1024


class DocumentExtractionError(ValueError):
    pass


class DocumentRunCancelled(Exception):
    pass


class DocumentRunTimedOut(Exception):
    pass


DOCUMENT_WORKER_ID = f"document-worker-{os.getenv('HOSTNAME', 'local')}-{os.getpid()}"


async def _assert_run_active(db: AsyncSession, run: DocumentAnalysisRun) -> None:
    await db.refresh(run, attribute_names=["status", "cancel_requested_at"])
    if run.status == "cancelled" or run.cancel_requested_at:
        raise DocumentRunCancelled("文档合规检查已停止")


async def _touch_run(db: AsyncSession, run: DocumentAnalysisRun) -> None:
    await _assert_run_active(db, run)
    now = datetime.utcnow()
    run.heartbeat_at = now
    run.lease_expires_at = now + timedelta(minutes=settings.DOCUMENT_LEASE_MINUTES)
    await db.commit()


async def cancel_document_run(db: AsyncSession, run: DocumentAnalysisRun, reason: str = "") -> None:
    if run.status not in {"queued", "running"}:
        raise ValueError(f"文档分析状态为 {run.status}，不能停止")
    now = datetime.utcnow()
    progress = {
        "stage": "cancelled",
        "percent": int((run.progress or {}).get("percent") or 0),
        "message": reason or "用户已停止文档合规检查",
    }
    stopped = await db.execute(
        update(DocumentAnalysisRun)
        .where(
            DocumentAnalysisRun.id == run.id,
            DocumentAnalysisRun.status.in_(["queued", "running"]),
        )
        .values(
            status="cancelled",
            cancel_requested_at=now,
            completed_at=now,
            lease_owner=None,
            lease_expires_at=None,
            progress=progress,
        )
    )
    if stopped.rowcount != 1:
        await db.rollback()
        raise ValueError("文档分析已结束，不能再停止")
    if run.task_id:
        task = await db.get(TaskInstance, run.task_id)
        if task and task.status == "in_progress":
            await db.execute(
                update(TaskInstance)
                .where(TaskInstance.id == task.id, TaskInstance.status == "in_progress")
                .values(
                    status="failed",
                    completed_at=now,
                    cancel_requested_at=now,
                    lease_owner=None,
                    lease_expires_at=None,
                    result={
                        "type": "doc_review",
                        "status": "cancelled",
                        "run_id": run.id,
                        "error": reason or "用户已停止文档合规检查",
                    },
                )
            )
    await db.commit()
    await db.refresh(run)


def build_retest_comparison(previous: dict[str, Any], current: dict[str, Any]) -> dict[str, Any]:
    before = float(previous.get("coverage") or 0)
    after = float(current.get("coverage") or 0)
    previous_gaps = {item["uid"]: item for item in previous.get("gap_items", [])}
    current_gaps = {item["uid"]: item for item in current.get("gap_items", [])}
    previous_reliable = previous.get("status") != "unable"
    current_reliable = current.get("status") != "unable"

    if previous_reliable and current_reliable:
        fixed_ids = sorted(previous_gaps.keys() - current_gaps.keys())
        new_ids = sorted(current_gaps.keys() - previous_gaps.keys())
        remaining_ids = sorted(previous_gaps.keys() & current_gaps.keys())
        status = "improved" if after > before else ("regressed" if after < before else "unchanged")
        delta = round(after - before, 2)
    else:
        fixed_ids = []
        new_ids = []
        remaining_ids = sorted(previous_gaps) if previous_reliable else []
        status = "unable" if not current_reliable else "baseline_unavailable"
        delta = 0

    return {
        "previous_status": previous.get("status"),
        "current_status": current.get("status"),
        "previous_coverage": before,
        "current_coverage": after,
        "delta": delta,
        "status": status,
        "initial_gaps": [item["reason"] for item in previous_gaps.values()],
        "current_gaps": [item["reason"] for item in current_gaps.values()],
        "fixed_gaps": [previous_gaps[uid]["reason"] for uid in fixed_ids],
        "remaining_gaps": [previous_gaps[uid]["reason"] for uid in remaining_ids],
        "new_gaps": [current_gaps[uid]["reason"] for uid in new_ids],
        "fixed_gap_ids": fixed_ids,
        "remaining_gap_ids": remaining_ids,
        "new_gap_ids": new_ids,
        "comparison_reliable": previous_reliable and current_reliable,
    }


def expand_document_upload(file_name: str, content: bytes) -> tuple[list[tuple[str, bytes]], list[str]]:
    """Return upload documents, safely expanding ZIP, RAR or 7z archives."""
    suffix = Path(file_name).suffix.lower()
    if suffix not in ARCHIVE_SUFFIXES:
        return [(safe_document_name(file_name), content)], []
    if suffix != ".zip":
        return _expand_libarchive(file_name, content)
    try:
        archive = zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile as exc:
        raise DocumentExtractionError("ZIP 压缩包已损坏或格式无效。") from exc

    documents: list[tuple[str, bytes]] = []
    skipped: list[str] = []
    total_size = 0
    with archive:
        for item in archive.infolist():
            path = Path(item.filename)
            if item.is_dir() or "__MACOSX" in path.parts or path.name.startswith("."):
                continue
            if path.is_absolute() or ".." in path.parts:
                raise DocumentExtractionError("ZIP 包含不安全的文件路径，已拒绝处理。")
            if item.flag_bits & 0x1:
                raise DocumentExtractionError("ZIP 包含加密文件，暂不支持处理。")
            suffix = path.suffix.lower()
            if suffix in ARCHIVE_SUFFIXES:
                raise DocumentExtractionError("暂不支持嵌套压缩包，请先解压后重新打包。")
            if suffix not in SUPPORTED_SUFFIXES:
                skipped.append(item.filename)
                continue
            total_size += item.file_size
            if item.file_size > 100 * 1024 * 1024 or total_size > MAX_BATCH_UNCOMPRESSED:
                raise DocumentExtractionError("ZIP 解压后超过单文件 100MB 或总计 300MB 限制。")
            documents.append((path.as_posix(), archive.read(item)))
            if len(documents) > MAX_BATCH_FILES:
                raise DocumentExtractionError(f"单次最多处理 {MAX_BATCH_FILES} 个文档。")
    if not documents:
        raise DocumentExtractionError("ZIP 中未发现支持的文档。")
    return documents, skipped


def _expand_libarchive(file_name: str, content: bytes) -> tuple[list[tuple[str, bytes]], list[str]]:
    binary = shutil.which("bsdtar")
    if not binary:
        raise DocumentExtractionError("RAR/7z 解压组件不可用，请检查后端镜像中的 libarchive-tools。")
    archive_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=Path(file_name).suffix, delete=False) as archive_file:
            archive_file.write(content)
            archive_path = Path(archive_file.name)
        listing = subprocess.run(
            [binary, "-tf", str(archive_path)],
            capture_output=True,
            text=True,
            timeout=30,
            check=False,
        )
        if listing.returncode:
            raise DocumentExtractionError("压缩包已损坏、已加密或格式不受支持。")
        documents: list[tuple[str, bytes]] = []
        skipped: list[str] = []
        total_size = 0
        for member in listing.stdout.splitlines():
            path = Path(member)
            if not member or member.endswith("/") or "__MACOSX" in path.parts or path.name.startswith("."):
                continue
            if path.is_absolute() or ".." in path.parts:
                raise DocumentExtractionError("压缩包包含不安全的文件路径，已拒绝处理。")
            if path.suffix.lower() in ARCHIVE_SUFFIXES:
                raise DocumentExtractionError("暂不支持嵌套压缩包，请先解压后重新打包。")
            if path.suffix.lower() not in SUPPORTED_SUFFIXES:
                skipped.append(member)
                continue
            process = subprocess.Popen(
                [binary, "-xOf", str(archive_path), "--", member],
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
            )
            chunks: list[bytes] = []
            member_size = 0
            assert process.stdout is not None
            while chunk := process.stdout.read(1024 * 1024):
                member_size += len(chunk)
                total_size += len(chunk)
                if member_size > 100 * 1024 * 1024 or total_size > MAX_BATCH_UNCOMPRESSED:
                    process.kill()
                    process.wait()
                    raise DocumentExtractionError("压缩包解压后超过单文件 100MB 或总计 300MB 限制。")
                chunks.append(chunk)
            if process.wait(timeout=30):
                raise DocumentExtractionError(f"无法解压 {path.name}，压缩包可能已加密或损坏。")
            documents.append((path.as_posix(), b"".join(chunks)))
            if len(documents) > MAX_BATCH_FILES:
                raise DocumentExtractionError(f"单次最多处理 {MAX_BATCH_FILES} 个文档。")
        if not documents:
            raise DocumentExtractionError("压缩包中未发现支持的文档。")
        return documents, skipped
    except subprocess.TimeoutExpired as exc:
        raise DocumentExtractionError("压缩包解析超时。") from exc
    finally:
        if archive_path:
            archive_path.unlink(missing_ok=True)


def safe_document_name(value: str) -> str:
    path = Path((value or "document").replace("\\", "/"))
    if path.is_absolute() or ".." in path.parts:
        raise DocumentExtractionError("文件路径不安全，已拒绝处理。")
    parts = [part for part in path.parts if part not in {"", "."}]
    return "/".join(parts) or "document"


def _block(
    document_file: DocumentFile,
    index: int,
    text: str,
    block_type: str = "text",
    page: int | None = None,
    section: str | None = None,
    source: str = "native",
    confidence: float = 1.0,
    table: list[list[str]] | None = None,
    bbox: list[float] | None = None,
) -> dict[str, Any]:
    normalized = re.sub(r"\s+", " ", text or "").strip()
    return {
        "block_id": f"f{document_file.id}-b{index}",
        "document_file_id": document_file.id,
        "file_name": document_file.original_name,
        "page": page,
        "section": section,
        "type": block_type,
        "bbox": bbox,
        "text": normalized,
        "table": table,
        "source": source,
        "confidence": confidence,
        "content_hash": hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
    }


def _docx_native(path: Path, document_file: DocumentFile) -> tuple[list[dict], list[tuple[str, bytes]], int]:
    import docx
    from docx.table import Table

    with zipfile.ZipFile(path) as archive:
        if sum(item.file_size for item in archive.infolist()) > 300 * 1024 * 1024:
            raise DocumentExtractionError("DOCX 解压后内容超过 300MB，已拒绝处理。")

    document = docx.Document(str(path))
    blocks: list[dict] = []
    current_section = None

    for item in document.iter_inner_content():
        if isinstance(item, Table):
            rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            text = "\n".join(" | ".join(row) for row in rows if any(row))
            if text.strip():
                blocks.append(_block(document_file, len(blocks), text, "table", section=current_section, table=rows))
            continue
        text = item.text.strip()
        if not text:
            continue
        style = (item.style.name or "").lower() if item.style else ""
        block_type = "heading" if "heading" in style or "标题" in style else "text"
        if block_type == "heading":
            current_section = text
        blocks.append(_block(document_file, len(blocks), text, block_type, section=current_section))

    for section in document.sections:
        for block_type, container in (("header", section.header), ("footer", section.footer)):
            text = "\n".join(p.text.strip() for p in container.paragraphs if p.text.strip())
            if text:
                blocks.append(_block(document_file, len(blocks), text, block_type))

    images: list[tuple[str, bytes]] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/") and Path(name).suffix.lower() in IMAGE_SUFFIXES:
                images.append((Path(name).name, archive.read(name)))
    return blocks, images, max(1, len(document.sections))


def _pdf_native(path: Path, document_file: DocumentFile, analysis_mode: str = "standard") -> tuple[list[dict], list[tuple[str, bytes]], int]:
    import pypdf
    import pypdfium2 as pdfium

    reader = pypdf.PdfReader(str(path))
    blocks: list[dict] = []
    visual_pages: list[tuple[str, bytes]] = []
    pdf = pdfium.PdfDocument(str(path))
    for page_index, page in enumerate(reader.pages):
        fragments: list[tuple[float, float, float, str]] = []

        def visitor(text, cm, tm, _font, font_size):
            value = re.sub(r"\s+", " ", text or "").strip()
            if value:
                fragments.append((float(tm[4] + cm[4]), float(tm[5] + cm[5]), float(font_size or 10), value))

        text = (page.extract_text(visitor_text=visitor) or "").strip()
        if fragments:
            lines: dict[int, list[tuple[float, float, str]]] = {}
            for x, y, size, value in fragments:
                lines.setdefault(round(y / 4), []).append((x, size, value))
            for line_key in sorted(lines, reverse=True):
                line = sorted(lines[line_key], key=lambda item: item[0])
                line_text = " ".join(item[2] for item in line).strip()
                if not line_text:
                    continue
                min_x = line[0][0]
                max_size = max(item[1] for item in line)
                max_x = max(item[0] + len(item[2]) * item[1] * 0.55 for item in line)
                y = line_key * 4
                blocks.append(_block(
                    document_file,
                    len(blocks),
                    line_text,
                    "text",
                    page=page_index + 1,
                    bbox=[min_x, y, max_x, y + max_size],
                ))
        elif text:
            blocks.append(_block(document_file, len(blocks), text, "text", page=page_index + 1))
        resources = page.get("/Resources") or {}
        has_images = bool(resources.get("/XObject"))
        if analysis_mode == "deep" or len(text) < 80 or has_images:
            bitmap = pdf[page_index].render(scale=1.5)
            buffer = io.BytesIO()
            bitmap.to_pil().save(buffer, format="PNG")
            visual_pages.append((f"page-{page_index + 1}.png", buffer.getvalue()))
    return blocks, visual_pages, len(reader.pages)


def _text_native(path: Path, document_file: DocumentFile) -> tuple[list[dict], list[tuple[str, bytes]], int]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = [
        _block(document_file, index, part, "text")
        for index, part in enumerate(re.split(r"\n\s*\n", text))
        if part.strip()
    ]
    return blocks, [], 1


async def _vision_blocks(
    document_file: DocumentFile,
    images: list[tuple[str, bytes]],
    start: int,
    analysis_mode: str,
    native_available: bool,
) -> tuple[list[dict], list[str], int]:
    blocks: list[dict] = []
    warnings: list[str] = []
    successful_images = 0
    async with httpx.AsyncClient(timeout=180) as client:
        for image_name, image_bytes in images:
            try:
                page_match = re.search(r"page-(\d+)", image_name)
                default_page = int(page_match.group(1)) if page_match else None
                response = await client.post(
                    f"{settings.OCR_SERVER_URL}/execute",
                    json={
                        "tool": "document_page_parse",
                        "params": {
                            "image_base64": base64.b64encode(image_bytes).decode(),
                            "file_name": image_name,
                            "use_chart_recognition": True,
                            "analysis_mode": analysis_mode,
                            "cross_validate": analysis_mode == "deep",
                            "prefer_vision": analysis_mode == "deep" or not native_available,
                        },
                    },
                )
                response.raise_for_status()
                result = response.json()
                if result.get("status") != "success":
                    raise DocumentExtractionError(result.get("error") or "视觉解析服务未返回成功状态。")
                payload = result.get("data", {})
                warnings.extend(f"{image_name}: {warning}" for warning in payload.get("warnings") or [] if str(warning).strip())
                image_block_count = 0
                for item in payload.get("blocks") or []:
                    text = item.get("text") or ""
                    if not text.strip():
                        continue
                    blocks.append(_block(
                        document_file,
                        start + len(blocks),
                        text,
                        block_type=item.get("type") or "image_text",
                        page=item.get("page") or default_page,
                        source=item.get("source") or "vision",
                        confidence=float(item.get("confidence") or 0.7),
                        table=item.get("table"),
                        bbox=item.get("bbox"),
                    ))
                    image_block_count += 1
                if image_block_count:
                    successful_images += 1
                else:
                    warnings.append(f"{image_name}: OCR/视觉解析未提取到内容")
            except Exception as exc:
                detail = str(exc).strip() or exc.__class__.__name__
                warnings.append(f"{image_name}: {detail}")
    return blocks, warnings, successful_images


def _deduplicate(blocks: list[dict]) -> list[dict]:
    result: list[dict] = []
    for block in sorted(blocks, key=lambda item: 0 if item["source"] == "native" else 1):
        block = dict(block)
        text_key = re.sub(r"[\W_]+", "", block.get("text", "")).lower()
        if not text_key:
            continue
        duplicate = next((item for item in result if (
            item.get("document_file_id") == block.get("document_file_id")
            and item.get("page") == block.get("page")
            and SequenceMatcher(
                None,
                re.sub(r"[\W_]+", "", item.get("text", "")).lower(),
                text_key,
            ).ratio() >= 0.94
        )), None)
        if duplicate:
            metadata = duplicate.setdefault("metadata", {})
            metadata.setdefault("corroborated_by", []).append({
                "source": block.get("source"),
                "confidence": block.get("confidence"),
                "bbox": block.get("bbox"),
            })
            duplicate["confidence"] = max(float(duplicate.get("confidence") or 0), float(block.get("confidence") or 0))
            continue
        result.append(block)
    for index, block in enumerate(result):
        block["block_id"] = f"f{block.get('document_file_id')}-b{index}"
    return result


def normalize_analysis_mode(mode: str | None) -> str:
    return mode if mode in ANALYSIS_MODES else "standard"


def _visual_coverage_state(
    native: list[dict],
    image_count: int,
    successful_images: int,
    warnings: list[str],
) -> dict[str, bool]:
    native_chars = len(re.sub(r"\s+", "", "".join(str(block.get("text") or "") for block in native)))
    visual_required = image_count > 0 and native_chars < 80
    return {
        "visual_required": visual_required,
        "visual_incomplete": visual_required and successful_images < image_count,
        "visual_degraded": bool(warnings) or successful_images < image_count,
    }


async def extract_document(document_file: DocumentFile, analysis_mode: str = "standard") -> dict[str, Any]:
    analysis_mode = normalize_analysis_mode(analysis_mode)
    path = file_storage.base_path / str(document_file.storage_path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        if suffix == ".doc":
            raise DocumentExtractionError("暂不支持旧版 DOC，请转换为 DOCX 或 PDF。")
        raise DocumentExtractionError(f"不支持的文件格式：{suffix or '未知'}")
    if not path.exists():
        raise DocumentExtractionError("上传文件不存在或存储不可用。")

    try:
        if suffix == ".docx":
            native, images, pages = _docx_native(path, document_file)
        elif suffix == ".pdf":
            native, images, pages = _pdf_native(path, document_file, analysis_mode)
        elif suffix in IMAGE_SUFFIXES:
            native, images, pages = [], [(document_file.original_name or path.name, path.read_bytes())], 1
        else:
            native, images, pages = _text_native(path, document_file)
    except Exception as exc:
        raise DocumentExtractionError(f"文件解析失败：{exc}") from exc

    vision, warnings, successful_images = await _vision_blocks(
        document_file,
        images,
        len(native),
        analysis_mode,
        bool(native),
    ) if images else ([], [], 0)
    blocks = _deduplicate([*native, *vision])
    if not blocks:
        detail = f"；视觉解析失败：{'；'.join(warnings)}" if warnings else ""
        raise DocumentExtractionError(f"未提取到可分析的文档内容{detail}")
    visual_state = _visual_coverage_state(native, len(images), successful_images, warnings)
    return {
        "blocks": blocks,
        "analysis_mode": analysis_mode,
        "page_count": pages,
        "native_blocks": len(native),
        "ocr_blocks": sum(1 for block in vision if block.get("source") == "ocr"),
        "vision_blocks": len(vision),
        "warnings": warnings,
        "visual_images": len(images),
        "visual_images_succeeded": successful_images,
        **visual_state,
    }


async def create_document_run(
    db: AsyncSession,
    task: TaskInstance,
    project_id: int,
    user_id: int,
    analysis_mode: str | None = None,
    run_parameters: dict[str, Any] | None = None,
) -> DocumentAnalysisRun:
    mode = normalize_analysis_mode(analysis_mode)
    phase = await db.get(PhaseInstance, task.phase_id)
    assessment = await db.get(Assessment, phase.assessment_id) if phase else None
    if not phase or not assessment or assessment.project_id != project_id:
        raise DocumentExtractionError("文档任务不属于有效的测评流程。")
    files = (await db.execute(
        select(DocumentFile).where(
            DocumentFile.assessment_id == assessment.id,
            DocumentFile.task_id == task.id,
            DocumentFile.is_active.is_(True),
        ).order_by(DocumentFile.created_at)
    )).scalars().all()
    if not files:
        raise DocumentExtractionError("该任务尚未上传或归类文档。")
    from app.services.report_service import invalidate_report_artifacts
    await invalidate_report_artifacts(db, project_id, "已重新分析文档材料")
    previous_run_id = (await db.execute(
        select(DocumentAnalysisRun.id).where(
            DocumentAnalysisRun.task_id == task.id,
            DocumentAnalysisRun.status == "completed",
        ).order_by(DocumentAnalysisRun.created_at.desc()).limit(1)
    )).scalar_one_or_none()
    run = DocumentAnalysisRun(
        project_id=project_id,
        assessment_id=assessment.id,
        phase_id=phase.id,
        task_id=task.id,
        requested_by=user_id,
        run_kind="retest" if previous_run_id else "initial",
        analysis_mode=mode,
        parameters={"previous_run_id": previous_run_id, **(run_parameters or {})},
        status="queued",
        progress={"stage": "queued", "percent": 0, "message": "等待文档分析"},
    )
    db.add(run)
    await db.flush()
    db.add_all(DocumentRunFile(analysis_run_id=run.id, document_file_id=file.id) for file in files)
    if not (run.parameters or {}).get("verification_run_id"):
        task.status = "in_progress"
        task.result = {"type": "doc_review", "status": "queued", "run_id": run.id, "analysis_mode": mode, "progress": run.progress}
    await db.commit()
    return run


async def create_document_batch_run(
    db: AsyncSession,
    phase_id: int,
    project_id: int,
    document_file_ids: list[int],
    user_id: int,
    analysis_mode: str | None = None,
    skipped_files: list[str] | None = None,
    duplicate_files: list[dict] | None = None,
    run_parameters: dict[str, Any] | None = None,
) -> DocumentAnalysisRun:
    mode = normalize_analysis_mode(analysis_mode)
    phase = await db.get(PhaseInstance, phase_id)
    assessment = await db.get(Assessment, phase.assessment_id) if phase else None
    if not phase or not assessment or assessment.project_id != project_id:
        raise DocumentExtractionError("批量文档不属于有效的测评流程。")
    from app.services.report_service import invalidate_report_artifacts
    await invalidate_report_artifacts(db, project_id, "已上传并分析新的批量文档")
    run = DocumentAnalysisRun(
        project_id=project_id,
        assessment_id=assessment.id,
        phase_id=phase.id,
        task_id=None,
        requested_by=user_id,
        run_kind="batch",
        analysis_mode=mode,
        status="queued",
        parameters={
            "skipped_files": skipped_files or [],
            "duplicate_files": duplicate_files or [],
            **(run_parameters or {}),
        },
        progress={"stage": "queued", "percent": 0, "message": "等待批量文档归类"},
    )
    db.add(run)
    await db.flush()
    db.add_all(DocumentRunFile(analysis_run_id=run.id, document_file_id=file_id) for file_id in document_file_ids)
    await db.execute(update(DocumentFile).where(
        DocumentFile.id.in_(document_file_ids),
        DocumentFile.uploaded_in_run_id.is_(None),
    ).values(uploaded_in_run_id=run.id))
    await db.commit()
    return run


async def _replace_batch_document_versions(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    classified: list[dict],
) -> int:
    """Deactivate prior files only for document categories present in a remediation batch."""
    files_by_task: dict[int, list[int]] = {}
    for item in classified:
        files_by_task.setdefault(int(item["task_id"]), []).append(int(item["document_file_id"]))
    replaced = 0
    for task_id, current_ids in files_by_task.items():
        old_files = list((await db.execute(select(DocumentFile).where(
            DocumentFile.assessment_id == run.assessment_id,
            DocumentFile.task_id == task_id,
            DocumentFile.is_active.is_(True),
            DocumentFile.id.not_in(current_ids),
        ))).scalars().all())
        replacement_id = current_ids[0]
        for document in old_files:
            document.is_active = False
            document.replaced_by_id = replacement_id
            await db.execute(update(DocumentBlock).where(
                DocumentBlock.document_file_id == document.id
            ).values(is_active=False, embedding=None))
            await knowledge_graph.purge_file(db, document.id)
        replaced += len(old_files)
    return replaced


async def _files_for_run(db: AsyncSession, run_id: int) -> list[DocumentFile]:
    return list((await db.execute(
        select(DocumentFile)
        .join(DocumentRunFile, DocumentRunFile.document_file_id == DocumentFile.id)
        .where(DocumentRunFile.analysis_run_id == run_id, DocumentFile.is_active.is_(True))
        .order_by(DocumentFile.created_at)
    )).scalars().all())


async def _persist_blocks(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    document_file: DocumentFile,
    extraction: dict[str, Any],
) -> list[dict[str, Any]]:
    await db.execute(delete(DocumentBlock).where(
        DocumentBlock.analysis_run_id == run.id,
        DocumentBlock.document_file_id == document_file.id,
    ))
    await db.execute(update(DocumentBlock).where(
        DocumentBlock.document_file_id == document_file.id,
        DocumentBlock.analysis_run_id != run.id,
        DocumentBlock.is_active.is_(True),
    ).values(is_active=False, embedding=None))
    vectors = [block.get("embedding") for block in extraction["blocks"]]
    embedding_model = next((block.get("embedding_model") for block in extraction["blocks"] if block.get("embedding_model")), None)
    embedding_error = None
    if not vectors or any(vector is None for vector in vectors):
        try:
            from app.services.llm_service import llm_service

            vectors = []
            models = []
            for start in range(0, len(extraction["blocks"]), 32):
                batch = extraction["blocks"][start:start + 32]
                result = await llm_service.embed_with_fallback(
                    db,
                    [str(block.get("text") or "")[:4000] for block in batch],
                    settings.DOCUMENT_EMBEDDING_DIMENSION,
                    input_type="passage",
                )
                vectors.extend(result["embeddings"])
                models.append(result["model"])
            embedding_model = models[0] if models and len(set(models)) == 1 else ",".join(dict.fromkeys(models))
        except SQLAlchemyError:
            raise
        except Exception as exc:
            vectors = [None] * len(extraction["blocks"])
            embedding_error = str(exc)

    rows = []
    for ordinal, block in enumerate(extraction["blocks"]):
        row = DocumentBlock(
            project_id=run.project_id,
            assessment_id=run.assessment_id,
            analysis_run_id=run.id,
            document_file_id=document_file.id,
            ordinal=ordinal,
            page_number=block.get("page"),
            section_path=[block["section"]] if block.get("section") else [],
            block_type=block.get("type") or "text",
            source=block.get("source") or "native",
            source_confidence=float(block.get("confidence") or 0),
            bbox=block.get("bbox"),
            text=block.get("text") or "",
            table_data=block.get("table"),
            content_sha256=block["content_hash"],
            metadata_json=block.get("metadata") or {},
            embedding_model=embedding_model if vectors[ordinal] is not None else None,
            embedding=vectors[ordinal],
            is_active=True,
        )
        db.add(row)
        rows.append(row)
    await db.flush()
    document_file.page_count = extraction["page_count"]
    document_file.parse_status = "parsed"
    document_file.extraction_summary = {
        key: extraction[key]
        for key in (
            "analysis_mode", "page_count", "native_blocks", "ocr_blocks", "vision_blocks", "warnings",
            "visual_images", "visual_images_succeeded", "visual_required", "visual_incomplete", "visual_degraded",
        )
    }
    document_file.extraction_summary.update({
        "embedding_status": "ready" if embedding_model and not embedding_error else "unavailable",
        "embedding_model": embedding_model,
        "embedding_error": embedding_error,
    })
    await knowledge_graph.sync_document_structure(
        db,
        project_id=run.project_id,
        assessment_id=run.assessment_id,
        phase_id=run.phase_id,
        task_id=run.task_id,
        run_id=run.id,
        file_id=document_file.id,
        blocks=[{
            "id": row.id,
            "ordinal": row.ordinal,
            "page_number": row.page_number,
            "section_path": row.section_path,
            "block_type": row.block_type,
            "content_sha256": row.content_sha256,
        } for row in rows],
    )
    return [{
        "block_id": row.id,
        "document_file_id": document_file.id,
        "file_name": document_file.original_name,
        "page": row.page_number,
        "section": row.section_path[-1] if row.section_path else None,
        "type": row.block_type,
        "bbox": row.bbox,
        "text": row.text,
        "table": row.table_data,
        "source": row.source,
        "confidence": row.source_confidence,
        "content_hash": row.content_sha256,
        "metadata": row.metadata_json,
        "embedding_model": row.embedding_model,
        "embedding": list(row.embedding) if row.embedding is not None else None,
    } for row in rows]


async def _extract_and_store(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    document_file: DocumentFile,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    current_rows = (await db.execute(
        select(DocumentBlock).where(
            DocumentBlock.analysis_run_id == run.id,
            DocumentBlock.document_file_id == document_file.id,
        ).order_by(DocumentBlock.ordinal)
    )).scalars().all()
    if current_rows and document_file.parse_status == "parsed" and document_file.extraction_summary:
        summary = dict(document_file.extraction_summary)
        summary["cache_hit"] = True
        return [{
            "block_id": row.id,
            "document_file_id": document_file.id,
            "file_name": document_file.original_name,
            "page": row.page_number,
            "section": row.section_path[-1] if row.section_path else None,
            "type": row.block_type,
            "bbox": row.bbox,
            "text": row.text,
            "table": row.table_data,
            "source": row.source,
            "confidence": row.source_confidence,
            "content_hash": row.content_sha256,
            "metadata": row.metadata_json,
            "embedding_model": row.embedding_model,
            "embedding": list(row.embedding) if row.embedding is not None else None,
        } for row in current_rows], summary

    cached_run_id = (await db.execute(
        select(DocumentBlock.analysis_run_id)
        .join(DocumentAnalysisRun, DocumentAnalysisRun.id == DocumentBlock.analysis_run_id)
        .where(
            DocumentBlock.document_file_id == document_file.id,
            DocumentBlock.analysis_run_id != run.id,
            DocumentAnalysisRun.analysis_mode == run.analysis_mode,
            DocumentAnalysisRun.status.in_(["running", "completed"]),
        )
        .order_by(DocumentAnalysisRun.created_at.desc())
        .limit(1)
    )).scalar_one_or_none()
    if cached_run_id and document_file.extraction_summary:
        cached = (await db.execute(
            select(DocumentBlock).where(
                DocumentBlock.analysis_run_id == cached_run_id,
                DocumentBlock.document_file_id == document_file.id,
            ).order_by(DocumentBlock.ordinal)
        )).scalars().all()
        if cached:
            extraction = {
                **document_file.extraction_summary,
                "blocks": [{
                    "page": row.page_number,
                    "section": row.section_path[-1] if row.section_path else None,
                    "type": row.block_type,
                    "bbox": row.bbox,
                    "text": row.text,
                    "table": row.table_data,
                    "source": row.source,
                    "confidence": row.source_confidence,
                    "content_hash": row.content_sha256,
                    "metadata": row.metadata_json,
                    "embedding_model": row.embedding_model,
                    "embedding": list(row.embedding) if row.embedding is not None else None,
                } for row in cached],
                "cache_hit": True,
            }
            blocks = await _persist_blocks(db, run, document_file, extraction)
            summary = dict(document_file.extraction_summary or {})
            summary["cache_hit"] = True
            return blocks, summary
    extraction = await extract_document(document_file, run.analysis_mode)
    blocks = await _persist_blocks(db, run, document_file, extraction)
    summary = dict(document_file.extraction_summary or {})
    summary["cache_hit"] = False
    return blocks, summary


_TRANSIENT_DOCUMENT_ERRORS = (
    "readtimeout", "timeout", "timed out", "超时", "temporarily unavailable", "service unavailable",
    "connection reset", "connection refused", "server disconnected", "remote protocol error",
    "internal server error", "bad gateway", "gateway timeout", " 500", " 502", " 503", " 504",
)


def _is_transient_document_error(exc: Exception) -> bool:
    message = f"{exc.__class__.__name__}: {exc}".lower()
    return any(marker in message for marker in _TRANSIENT_DOCUMENT_ERRORS)


def _extraction_failure_message(document_file: DocumentFile, exc: Exception) -> str:
    attempts = int((document_file.extraction_summary or {}).get("attempts") or 1)
    return f"{exc}（已尝试 {attempts} 次）" if attempts > 1 else str(exc)


async def _extract_with_retry(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    document_file: DocumentFile,
) -> tuple[list[dict], dict]:
    attempts = max(1, settings.DOCUMENT_FILE_RETRY_ATTEMPTS)
    for attempt in range(1, attempts + 1):
        try:
            blocks, summary = await _extract_and_store(db, run, document_file)
            summary = {**summary, "attempts": attempt, "recovered_after_retry": attempt > 1}
            document_file.extraction_summary = summary
            return blocks, summary
        except (DocumentRunCancelled, SQLAlchemyError):
            raise
        except Exception as exc:
            retryable = _is_transient_document_error(exc)
            document_file.extraction_summary = {
                **(document_file.extraction_summary or {}),
                "analysis_mode": run.analysis_mode,
                "error": str(exc),
                "attempts": attempt,
                "retryable": retryable,
            }
            if not retryable or attempt >= attempts:
                raise
            run.progress = {
                "stage": "native_extraction",
                "percent": int((run.progress or {}).get("percent") or 5),
                "message": f"{document_file.original_name} 提取超时，正在进行第 {attempt + 1}/{attempts} 次尝试",
            }
            await _touch_run(db, run)
            await asyncio.sleep(min(2 ** (attempt - 1), 4))


async def _persist_control_results(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    analysis: dict[str, Any],
) -> None:
    await db.execute(delete(DocumentControlResult).where(DocumentControlResult.analysis_run_id == run.id))
    for control in analysis.get("controls") or []:
        failed = [point for point in control.get("points") or [] if point.get("status") in {"fail", "partial", "unable"}]
        reason = "；".join(
            point.get("llm_reason") or point.get("missing_judgement") or point.get("text") or "证据不足"
            for point in failed
        ) or "所有必需证据均已定位且未发现矛盾。"
        result = DocumentControlResult(
            project_id=run.project_id,
            assessment_id=run.assessment_id,
            task_id=run.task_id,
            analysis_run_id=run.id,
            control_uid=control.get("uid") or control.get("id"),
            verdict=control.get("status") or "unable",
            confidence=float(analysis.get("confidence") or 0),
            reason=reason,
            missing_requirements=[point.get("uid") for point in failed if point.get("status") == "fail"],
            contradictory_requirements=[point.get("uid") for point in failed if point.get("contradiction")],
            rule_snapshot={
                "control_id": control.get("id"),
                "title": control.get("title"),
                "points": [{key: point.get(key) for key in ("uid", "id", "text", "status", "severity", "missing_judgement")} for point in control.get("points") or []],
            },
            model_snapshot={
                "engine": analysis.get("evidence_engine"),
                "error": analysis.get("llm_review_error"),
                "points": [{
                    "uid": point.get("uid"),
                    "decision": point.get("status"),
                    "confidence": point.get("decision_confidence"),
                    "reason": point.get("llm_reason"),
                    "contradiction": bool(point.get("contradiction")),
                } for point in control.get("points") or []],
            },
        )
        db.add(result)
        await db.flush()
        for point in control.get("points") or []:
            requirement_uid = point.get("uid")
            if not requirement_uid:
                continue
            evidence = point.get("evidence") or []
            if not evidence and point.get("status") == "fail":
                await knowledge_graph.sync_missing_requirement(
                    db,
                    project_id=run.project_id,
                    assessment_id=run.assessment_id,
                    phase_id=run.phase_id,
                    task_id=run.task_id,
                    run_id=run.id,
                    result_id=result.id,
                    control_uid=result.control_uid,
                    requirement_uid=requirement_uid,
                )
            for rank, item in enumerate(evidence, start=1):
                stance = "contradict" if point.get("contradiction") else ("support" if point.get("status") == "pass" else "partial")
                link = DocumentEvidenceLink(
                    control_result_id=result.id,
                    document_block_id=int(item["block_id"]),
                    requirement_uid=requirement_uid,
                    stance=stance,
                    confidence=float(point.get("decision_confidence") or item.get("confidence") or analysis.get("confidence") or 0),
                    rationale=point.get("llm_reason") or f"命中：{', '.join(item.get('matched_keywords') or [])}",
                    rank=rank,
                )
                db.add(link)
                await knowledge_graph.sync_evidence_link(
                    db,
                    project_id=run.project_id,
                    assessment_id=run.assessment_id,
                    phase_id=run.phase_id,
                    task_id=run.task_id,
                    run_id=run.id,
                    result_id=result.id,
                    control_uid=result.control_uid,
                    requirement_uid=requirement_uid,
                    block_id=int(item["block_id"]),
                    stance=stance,
                    confidence=link.confidence,
                )


async def process_document_run(db: AsyncSession, run: DocumentAnalysisRun) -> None:
    await _assert_run_active(db, run)
    task = await db.get(TaskInstance, run.task_id) if run.task_id else None
    if not task:
        raise DocumentExtractionError("关联的文档检查任务不存在。")
    files = await _files_for_run(db, run.id)
    if not files:
        raise DocumentExtractionError("该任务尚未上传文档。")
    previous_run_id = (run.parameters or {}).get("previous_run_id")
    previous_run = await db.get(DocumentAnalysisRun, previous_run_id) if previous_run_id else None
    previous_analysis = previous_run.result_summary if previous_run else None
    analysis_mode = run.analysis_mode
    user_id = run.requested_by or 0
    is_verification = bool((run.parameters or {}).get("verification_run_id"))

    run.progress = {"stage": "native_extraction", "percent": 5, "message": "正在执行原生内容提取"}
    if not is_verification:
        task.result = {"type": "doc_review", "status": "processing", "run_id": run.id, "analysis_mode": analysis_mode, "progress": run.progress}
    await db.commit()

    all_blocks: list[dict] = []
    total_pages = 0
    manifests = []
    extraction_failures = []
    visual_incomplete = False
    for index, document_file in enumerate(files):
        await _touch_run(db, run)
        try:
            blocks, extraction_summary = await _extract_with_retry(db, run, document_file)
            total_pages += extraction_summary["page_count"]
            if total_pages > settings.DOCUMENT_MAX_TOTAL_PAGES:
                raise DocumentExtractionError(f"文档总页数超过 {settings.DOCUMENT_MAX_TOTAL_PAGES} 页限制。")
            all_blocks.extend(blocks)
            visual_incomplete = visual_incomplete or bool(extraction_summary.get("visual_incomplete"))
            manifests.append({
                "document_file_id": document_file.id,
                "file_name": document_file.original_name,
                "status": "parsed",
                **extraction_summary,
            })
        except (DocumentRunCancelled, SQLAlchemyError):
            raise
        except Exception as exc:
            error_message = _extraction_failure_message(document_file, exc)
            document_file.parse_status = "failed"
            document_file.extraction_summary = {
                **(document_file.extraction_summary or {}),
                "analysis_mode": analysis_mode,
                "error": str(exc),
                "display_error": error_message,
            }
            failure = {
                "document_file_id": document_file.id,
                "file_name": document_file.original_name,
                "status": "failed",
                "error": error_message,
                "attempts": (document_file.extraction_summary or {}).get("attempts", 1),
            }
            extraction_failures.append(failure)
            manifests.append(failure)
        run.progress = {
            "stage": "fusion",
            "percent": 10 + int(45 * (index + 1) / len(files)),
            "message": f"已提取并融合 {index + 1}/{len(files)} 个文件",
        }
        if not is_verification:
            task.result = {"type": "doc_review", "status": "processing", "run_id": run.id, "analysis_mode": analysis_mode, "progress": run.progress}
        await _touch_run(db, run)

    if not all_blocks:
        detail = "；".join(item["error"] for item in extraction_failures) or "没有可分析内容"
        raise DocumentExtractionError(f"全部文档均无法提取：{detail}")

    run.progress = {"stage": "retrieval", "percent": 65, "message": "正在从标准图谱召回检查项和候选证据"}
    if not is_verification:
        task.result = {"type": "doc_review", "status": "processing", "run_id": run.id, "analysis_mode": analysis_mode, "progress": run.progress}
    await _touch_run(db, run)
    expected_name = task.name.split("文档检查：", 1)[1].strip() if "文档检查：" in task.name else task.name
    control_engine = await DocumentControlEngine.from_graph(db)
    analysis = await control_engine.analyze_retrieved(db, run.id, expected_doc_name=expected_name)
    analysis["files"] = manifests
    analysis["analysis_mode"] = analysis_mode
    analysis["document_file_ids"] = [document_file.id for document_file in files]
    analysis["run_id"] = run.id
    analysis["extraction_failures"] = extraction_failures
    run.progress = {"stage": "judging", "percent": 82, "message": "正在执行结构化证据判定"}
    if not is_verification:
        task.result = {"type": "doc_review", "status": "processing", "run_id": run.id, "analysis_mode": analysis_mode, "progress": run.progress}
    await _touch_run(db, run)
    analysis = await control_engine.review_with_llm(db, user_id, analysis)
    if extraction_failures or visual_incomplete:
        reasons = []
        if extraction_failures:
            reasons.append(f"{len(extraction_failures)} 个文件提取失败")
        if visual_incomplete:
            reasons.append("必要的视觉解析未完整完成")
        analysis["status"] = "unable"
        analysis["confidence"] = 0
        analysis["message"] = "；".join(reasons) + "，本次不能生成完整合规结论。"
    if previous_analysis and previous_analysis.get("type") == DOCUMENT_SOURCE:
        analysis["retest_comparison"] = build_retest_comparison(previous_analysis, analysis)

    run.progress = {"stage": "generating_results", "percent": 94, "message": "正在生成差距、整改和复测结果"}
    if not is_verification:
        task.result = {"type": "doc_review", "status": "processing", "run_id": run.id, "analysis_mode": analysis_mode, "progress": run.progress}
    await _touch_run(db, run)
    await _persist_control_results(db, run, analysis)
    from app.api.assessments import _sync_document_gap_findings
    sync = await _sync_document_gap_findings(db, run.project_id, task, analysis, user_id)
    analysis["gap_sync"] = sync
    unable = analysis.get("status") == "unable"
    task_result = {
        "type": "doc_review",
        "status": "unable" if unable else "completed",
        "analysis": analysis,
        "document_file_ids": analysis["document_file_ids"],
        "run_id": run.id,
        **({"error": analysis.get("message") or "文档合规分析无法形成可靠结论"} if unable else {}),
    }
    from app.services.flow_engine import get_flow_engine
    await _assert_run_active(db, run)
    if not is_verification:
        if unable:
            task.status = "failed"
            task.completed_at = datetime.utcnow()
            task.result = task_result
            await db.commit()
            await get_flow_engine(db).reconcile_phase_progress(task.phase_id)
        else:
            await get_flow_engine(db).complete_task(task.id, task_result)
    else:
        task.status = "failed" if unable else "completed"
        task.result = task_result
        task.completed_at = datetime.utcnow()
    run.status = "completed"
    run.result_summary = analysis
    if unable:
        run.error_code = None
        run.error_message = None
        run.progress = {
            "stage": "completed",
            "percent": 100,
            "message": analysis.get("message") or "分析完成，存在无法判断的检查项",
        }
    else:
        run.progress = {"stage": "completed", "percent": 100, "message": "文档合规分析完成"}
    run.completed_at = datetime.utcnow()
    run.lease_owner = None
    run.lease_expires_at = None
    await db.commit()


async def _record_batch_document_unable(
    db: AsyncSession,
    run: DocumentAnalysisRun,
    task: TaskInstance,
    message: str,
    files: list[dict[str, Any]],
) -> None:
    """Make a missing or unreadable required document visible to findings and reports."""
    analysis = {
        "type": DOCUMENT_SOURCE,
        "status": "unable",
        "document_name": task.name.removeprefix("文档检查："),
        "coverage": 0,
        "confidence": 0,
        "controls": [],
        "gaps": [message],
        "files": files,
        "analysis_mode": run.analysis_mode,
        "document_file_ids": [item["document_file_id"] for item in files if item.get("document_file_id")],
        "run_id": run.id,
        "message": message,
    }
    from app.api.assessments import _sync_document_gap_findings
    analysis["gap_sync"] = await _sync_document_gap_findings(
        db, run.project_id, task, analysis, run.requested_by or 0,
    )
    task.status = "failed"
    task.completed_at = datetime.utcnow()
    task.result = {
        "type": "doc_review",
        "status": "unable",
        "analysis": analysis,
        "document_file_ids": analysis["document_file_ids"],
        "run_id": run.id,
        "error": message,
    }


async def process_document_batch_run(db: AsyncSession, run: DocumentAnalysisRun) -> None:
    await _assert_run_active(db, run)
    parameters = run.parameters or {}
    verification_batch = bool(parameters.get("verification_batch"))
    task_phase_id = int(parameters.get("document_task_phase_id") or run.phase_id)
    tasks = (await db.execute(
        select(TaskInstance).where(TaskInstance.phase_id == task_phase_id, TaskInstance.task_type == "doc_review")
    )).scalars().all()
    task_by_name = {
        task.name.split("文档检查：", 1)[1].strip(): task
        for task in tasks
        if "文档检查：" in task.name
    }
    if not task_by_name:
        raise DocumentExtractionError("当前阶段没有可用的文档检查任务。")
    files = await _files_for_run(db, run.id)
    if not files:
        raise DocumentExtractionError("批量上传任务没有可分析的文档。")
    run.progress = {"stage": "native_extraction", "percent": 5, "message": "正在提取并识别文档"}
    await db.commit()

    engine = await DocumentControlEngine.from_graph(db)
    classified: list[dict] = []
    unclassified: list[dict] = []
    affected_task_ids: set[int] = set()
    unable_tasks: dict[int, dict[str, Any]] = {}
    total_pages = 0
    for index, document_file in enumerate(files):
        run.progress = {
            "stage": "native_extraction",
            "percent": 5 + int(5 * index / len(files)),
            "message": f"正在处理 {index + 1}/{len(files)}：{document_file.original_name}",
        }
        await _touch_run(db, run)
        try:
            blocks, extraction_summary = await _extract_with_retry(db, run, document_file)
            total_pages += extraction_summary.get("page_count") or 0
            if total_pages > settings.DOCUMENT_MAX_TOTAL_PAGES:
                raise DocumentExtractionError(f"文档总页数超过 {settings.DOCUMENT_MAX_TOTAL_PAGES} 页限制。")
            rule_classification = engine.classify_blocks(document_file.original_name, blocks)
            classification = await engine.classify_with_llm(
                db,
                run.requested_by or 0,
                document_file.original_name,
                blocks,
                rule_classification,
            )
            document_file.classification = classification
            matched_task = task_by_name.get(classification.get("document_name"))
            if classification["status"] == "classified" and matched_task:
                document_file.task_id = matched_task.id
                affected_task_ids.add(matched_task.id)
                classified.append({
                    "document_file_id": document_file.id,
                    "file_name": document_file.original_name,
                    "task_id": matched_task.id,
                    **classification,
                })
            else:
                unclassified.append({"document_file_id": document_file.id, "file_name": document_file.original_name, **classification})
        except (DocumentRunCancelled, SQLAlchemyError):
            raise
        except Exception as exc:
            error_message = _extraction_failure_message(document_file, exc)
            document_file.parse_status = "failed"
            document_file.extraction_summary = {
                **(document_file.extraction_summary or {}),
                "analysis_mode": run.analysis_mode,
                "error": str(exc),
                "display_error": error_message,
            }
            fallback = engine.classify_blocks(document_file.original_name, [])
            matched_task = task_by_name.get(fallback.get("document_name"))
            if fallback.get("status") == "classified" and matched_task:
                classification = {
                    **fallback,
                    "extraction_status": "unable",
                    "extraction_error": error_message,
                    "extraction_attempts": (document_file.extraction_summary or {}).get("attempts", 1),
                    "reason": f"{fallback.get('reason')}，但内容提取失败：{error_message}",
                }
                document_file.task_id = matched_task.id
                document_file.classification = classification
                classified.append({
                    "document_file_id": document_file.id,
                    "file_name": document_file.original_name,
                    "task_id": matched_task.id,
                    **classification,
                })
                if verification_batch:
                    affected_task_ids.add(matched_task.id)
                else:
                    unable_tasks[matched_task.id] = {
                        "message": f"已按文件名识别为“{fallback.get('document_name')}”，但未能提取可分析正文：{error_message}",
                        "files": [{
                            "document_file_id": document_file.id,
                            "file_name": document_file.original_name,
                            "status": "failed",
                            "error": error_message,
                            "attempts": (document_file.extraction_summary or {}).get("attempts", 1),
                        }],
                    }
            else:
                unclassified.append({
                    "document_file_id": document_file.id,
                    "file_name": document_file.original_name,
                    "status": "unclassified",
                    "naming_status": "unable",
                    "reason": error_message,
                    "confidence": 0,
                })
        run.progress = {
            "stage": "classification",
            "percent": 10 + int(60 * (index + 1) / len(files)),
            "message": f"已识别 {index + 1}/{len(files)} 个文档",
        }
        await _touch_run(db, run)

    replaced_files = await _replace_batch_document_versions(db, run, classified) if verification_batch else 0
    missing = []
    for document_name, task in task_by_name.items():
        has_document = (await db.execute(
            select(DocumentFile.id).where(
                DocumentFile.assessment_id == run.assessment_id,
                DocumentFile.task_id == task.id,
                DocumentFile.is_active.is_(True),
            ).limit(1)
        )).scalar_one_or_none()
        if not has_document:
            missing.append({"task_id": task.id, "document_name": document_name})

    if not verification_batch:
        for item in missing:
            unable_tasks.setdefault(item["task_id"], {
                "message": f"本次材料包未提供或未能可靠归类“{item['document_name']}”，无法完成该项合规检查。",
                "files": [],
            })
        for task_id, detail in unable_tasks.items():
            task = await db.get(TaskInstance, task_id)
            if task:
                await _record_batch_document_unable(db, run, task, detail["message"], detail["files"])

    run.result_summary = {
        "type": BATCH_DOCUMENT_SOURCE,
        "analysis_mode": run.analysis_mode,
        "total_files": len(files),
        "classified": classified,
        "unclassified": unclassified,
        "missing": missing,
        "skipped_files": parameters.get("skipped_files") or [],
        "duplicate_files": parameters.get("duplicate_files") or [],
        "coverage": round((len(task_by_name) - len(missing)) / len(task_by_name), 2),
        "verification_batch": verification_batch,
        "replaced_files": replaced_files,
        "verification_runs": [],
        "verification_skipped": [],
    }
    run.progress = {
        "stage": "analyzing",
        "percent": 85,
        "message": "归类完成，正在提交整改复测" if verification_batch else "归类完成，正在逐项执行合规分析",
    }
    await _touch_run(db, run)

    for task_id in sorted(affected_task_ids):
        await _touch_run(db, run)
        task = await db.get(TaskInstance, task_id)
        if task:
            if verification_batch:
                try:
                    from app.services.verification_service import queue_document_task_verification
                    verification, document_run, _, count = await queue_document_task_verification(
                        db,
                        project_id=run.project_id,
                        task=task,
                        actor_id=run.requested_by or 0,
                        analysis_mode=run.analysis_mode,
                        notes="批量提交整改材料后自动重新检查。",
                    )
                    run.result_summary["verification_runs"].append({
                        "task_id": task.id,
                        "document_name": task.name.removeprefix("文档检查："),
                        "verification_run_id": verification.id,
                        "document_run_id": document_run.id,
                        "finding_count": count,
                    })
                except ValueError as exc:
                    run.result_summary["verification_skipped"].append({
                        "task_id": task.id,
                        "document_name": task.name.removeprefix("文档检查："),
                        "reason": str(exc),
                    })
            else:
                await create_document_run(db, task, run.project_id, run.requested_by or 0, run.analysis_mode)

    if unable_tasks:
        from app.services.flow_engine import get_flow_engine
        await get_flow_engine(db).reconcile_phase_progress(task_phase_id)

    queued_count = len(run.result_summary["verification_runs"])
    skipped_count = len(run.result_summary["verification_skipped"])
    if verification_batch:
        message = f"整改材料已归类，{queued_count} 类文档已进入重新检查"
        if skipped_count:
            message += f"，{skipped_count} 类未提交"
    else:
        message = "文档已归类，合规分析任务已提交"
    run.result_summary = {
        **(run.result_summary or {}),
        "verification_runs": list((run.result_summary or {}).get("verification_runs") or []),
        "verification_skipped": list((run.result_summary or {}).get("verification_skipped") or []),
    }
    run.progress = {"stage": "completed", "percent": 100, "message": message}
    run.status = "completed"
    run.completed_at = datetime.utcnow()
    run.lease_owner = None
    run.lease_expires_at = None
    await db.commit()


async def recover_incomplete_document_runs(db: AsyncSession) -> int:
    now = datetime.utcnow()
    runs = (await db.execute(
        select(DocumentAnalysisRun).where(
            DocumentAnalysisRun.status == "running",
            or_(DocumentAnalysisRun.lease_expires_at.is_(None), DocumentAnalysisRun.lease_expires_at < now),
        )
    )).scalars().all()
    for run in runs:
        if run.cancel_requested_at:
            await cancel_document_run(db, run)
            continue
        run.lease_owner = None
        run.lease_expires_at = None
        if run.attempt_count >= settings.DOCUMENT_MAX_RECOVERY_ATTEMPTS:
            run.status = "failed"
            run.completed_at = now
            run.error_code = "document_recovery_exhausted"
            run.error_message = "文档分析连续中断，已达到自动恢复次数上限"
            run.progress = {"stage": "failed", "percent": 100, "message": run.error_message}
        else:
            percent = int((run.progress or {}).get("percent") or 0)
            run.status = "queued"
            run.progress = {
                "stage": "queued",
                "percent": percent,
                "message": f"检测到 Worker 中断，将从已完成文件继续（第 {run.attempt_count + 1} 次尝试）",
            }
    if runs:
        await db.commit()
    return len(runs)


async def _monitor_document_run(run_id: int) -> str:
    while True:
        await asyncio.sleep(max(1, settings.TASK_HEARTBEAT_SECONDS))
        async with AsyncSessionLocal() as monitor_db:
            run = await monitor_db.get(DocumentAnalysisRun, run_id)
            if not run:
                return "missing"
            if run.status == "cancelled" or run.cancel_requested_at:
                return "cancelled"
            if run.status != "running":
                return "finished"
            if run.lease_owner != DOCUMENT_WORKER_ID:
                return "lease_lost"
            now = datetime.utcnow()
            run.heartbeat_at = now
            run.lease_expires_at = now + timedelta(minutes=settings.DOCUMENT_LEASE_MINUTES)
            await monitor_db.commit()


async def _run_document_with_controls(db: AsyncSession, run: DocumentAnalysisRun) -> None:
    process = asyncio.create_task(
        process_document_batch_run(db, run) if run.run_kind == "batch" else process_document_run(db, run)
    )
    monitor = asyncio.create_task(_monitor_document_run(run.id))
    done, _ = await asyncio.wait(
        {process, monitor},
        timeout=max(60, settings.DOCUMENT_RUN_TIMEOUT_MINUTES * 60),
        return_when=asyncio.FIRST_COMPLETED,
    )
    if process in done:
        monitor.cancel()
        with contextlib.suppress(asyncio.CancelledError):
            await monitor
        await process
        return
    if monitor in done and monitor.result() == "finished":
        await process
        return
    process.cancel()
    monitor.cancel()
    with contextlib.suppress(asyncio.CancelledError):
        await process
    with contextlib.suppress(asyncio.CancelledError):
        await monitor
    if monitor in done and monitor.result() == "cancelled":
        raise DocumentRunCancelled("用户已停止文档合规检查")
    if not done:
        raise DocumentRunTimedOut(f"文档分析超过 {settings.DOCUMENT_RUN_TIMEOUT_MINUTES} 分钟，已停止本次尝试")
    raise DocumentRunTimedOut("文档分析执行权已失效，将由 Worker 恢复")


async def process_pending_document_runs(db: AsyncSession, limit: int | None = None) -> int:
    await recover_incomplete_document_runs(db)
    runs = (await db.execute(
        select(DocumentAnalysisRun)
        .where(
            DocumentAnalysisRun.status == "queued",
            DocumentAnalysisRun.cancel_requested_at.is_(None),
            DocumentAnalysisRun.attempt_count < settings.DOCUMENT_MAX_RECOVERY_ATTEMPTS,
        )
        .order_by(DocumentAnalysisRun.created_at)
        .limit(limit or settings.DOCUMENT_WORKER_BATCH_SIZE)
    )).scalars().all()
    processed = 0
    for run in runs:
        run_id = run.id
        now = datetime.utcnow()
        claimed = (await db.execute(
            update(DocumentAnalysisRun)
            .where(
                DocumentAnalysisRun.id == run.id,
                DocumentAnalysisRun.status == "queued",
                DocumentAnalysisRun.cancel_requested_at.is_(None),
            )
            .values(
                status="running",
                started_at=run.started_at or now,
                lease_owner=DOCUMENT_WORKER_ID,
                lease_expires_at=now + timedelta(minutes=settings.DOCUMENT_LEASE_MINUTES),
                heartbeat_at=now,
                attempt_count=(run.attempt_count or 0) + 1,
            )
            .returning(DocumentAnalysisRun.id)
        )).scalar_one_or_none()
        await db.commit()
        if not claimed:
            continue
        await db.refresh(run)
        processed += 1
        try:
            await _run_document_with_controls(db, run)
        except DocumentRunCancelled:
            await db.rollback()
            run = await db.get(DocumentAnalysisRun, run_id)
            if run and run.status != "cancelled":
                await cancel_document_run(db, run)
        except DocumentRunTimedOut as exc:
            await db.rollback()
            logger.warning("Document analysis run %s interrupted: %s", run_id, exc)
            run = await db.get(DocumentAnalysisRun, run_id)
            if not run or run.status == "cancelled":
                continue
            run.lease_owner = None
            run.lease_expires_at = None
            if run.attempt_count < settings.DOCUMENT_MAX_RECOVERY_ATTEMPTS:
                run.status = "queued"
                run.progress = {
                    "stage": "queued",
                    "percent": int((run.progress or {}).get("percent") or 0),
                    "message": f"{exc}；将从已完成文件继续",
                }
            else:
                run.status = "failed"
                run.completed_at = datetime.utcnow()
                run.error_code = "document_recovery_exhausted"
                run.error_message = str(exc)
                run.progress = {"stage": "failed", "percent": 100, "message": str(exc)}
                task = await db.get(TaskInstance, run.task_id) if run.task_id else None
                if task and task.status == "in_progress" and not (run.parameters or {}).get("verification_run_id"):
                    task.status = "failed"
                    task.completed_at = datetime.utcnow()
                    task.result = {"type": "doc_review", "status": "unable", "error": str(exc), "run_id": run.id}
            await db.commit()
        except Exception as exc:
            await db.rollback()
            logger.exception("Document analysis run %s failed", run_id)
            run = await db.get(DocumentAnalysisRun, run_id)
            if not run:
                continue
            task = await db.get(TaskInstance, run.task_id) if run.task_id else None
            if task and not (run.parameters or {}).get("verification_run_id"):
                task.status = "failed"
                task.result = {"type": "doc_review", "status": "unable", "error": str(exc), "run_id": run.id}
            run.status = "failed"
            run.lease_owner = None
            run.lease_expires_at = None
            run.error_code = "document_extraction_failed" if isinstance(exc, DocumentExtractionError) else "document_analysis_failed"
            run.error_message = str(exc)
            run.progress = {"stage": "failed", "percent": 100, "message": str(exc)}
            run.completed_at = datetime.utcnow()
            await db.commit()
    return processed
